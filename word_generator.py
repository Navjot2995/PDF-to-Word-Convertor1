from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re

class WordGenerator:
    def __init__(self):
        self.document = None
    
    def create_document(self, text_content, output_path, preserve_formatting=True):
        """Create a Word document from extracted text"""
        try:
            # Create a new document
            self.document = Document()
            
            # Set document properties
            self.document.core_properties.title = "Converted PDF Document"
            self.document.core_properties.author = "PDF to Word Converter"
            
            # Process and add content
            if preserve_formatting:
                self._add_formatted_content(text_content)
            else:
                self._add_plain_content(text_content)
            
            # Save the document
            self.document.save(output_path)
            
        except Exception as e:
            raise Exception(f"Error creating Word document: {e}")
    
    def _add_formatted_content(self, text_content):
        """Add content with formatting preservation"""
        try:
            # Split content into sections/pages
            sections = re.split(r'--- Page \d+ ---', text_content)
            
            for i, section in enumerate(sections):
                if not section.strip():
                    continue
                
                # Add page break for new pages (except first)
                if i > 0:
                    self.document.add_page_break()
                
                # Process paragraphs
                paragraphs = section.split('\n\n')
                
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    
                    # Detect different types of content
                    if self._is_heading(para_text):
                        self._add_heading(para_text)
                    elif self._is_list_item(para_text):
                        self._add_list_item(para_text)
                    else:
                        self._add_paragraph(para_text)
                        
        except Exception as e:
            print(f"Error adding formatted content: {e}")
            # Fallback to plain content
            self._add_plain_content(text_content)
    
    def _add_plain_content(self, text_content):
        """Add content as plain text"""
        try:
            # Split into paragraphs
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    # Handle page separators
                    if '--- Page' in para_text and '---' in para_text:
                        self.document.add_page_break()
                        continue
                    
                    # Add regular paragraph
                    paragraph = self.document.add_paragraph(para_text)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
        except Exception as e:
            print(f"Error adding plain content: {e}")
            # Final fallback - add as single paragraph
            self.document.add_paragraph(text_content)
    
    def _is_heading(self, text):
        """Detect if text is likely a heading"""
        # Common heading patterns
        patterns = [
            r'^[A-Z][A-Z\s]{2,}$',  # ALL CAPS
            r'^\d+\.\s*[A-Z]',      # Numbered headings
            r'^[A-Z][^.!?]*$',      # Capitalized without ending punctuation
            r'^\s*[A-Z][^.!?]{10,50}$'  # Short capitalized text
        ]
        
        text = text.strip()
        
        # Check patterns
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        # Check if it's short and doesn't end with punctuation
        if len(text) < 100 and not text.endswith(('.', '!', '?', ';', ':')):
            words = text.split()
            if len(words) <= 10 and text[0].isupper():
                return True
        
        return False
    
    def _is_list_item(self, text):
        """Detect if text is a list item"""
        patterns = [
            r'^\s*[-•*]\s+',        # Bullet points
            r'^\s*\d+\.\s+',        # Numbered lists
            r'^\s*[a-zA-Z]\.\s+',   # Lettered lists
            r'^\s*\(\d+\)\s+',      # Parenthetical numbers
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _add_heading(self, text):
        """Add text as a heading"""
        try:
            # Determine heading level based on text characteristics
            if text.isupper():
                level = 1
            elif len(text.split()) <= 5:
                level = 2
            else:
                level = 3
            
            heading = self.document.add_heading(text, level=level)
            
        except Exception as e:
            print(f"Error adding heading: {e}")
            # Fallback to bold paragraph
            para = self.document.add_paragraph()
            run = para.add_run(text)
            run.bold = True
    
    def _add_list_item(self, text):
        """Add text as a list item"""
        try:
            # Remove list markers
            clean_text = re.sub(r'^\s*[-•*]\s+', '', text)
            clean_text = re.sub(r'^\s*\d+\.\s+', '', clean_text)
            clean_text = re.sub(r'^\s*[a-zA-Z]\.\s+', '', clean_text)
            clean_text = re.sub(r'^\s*\(\d+\)\s+', '', clean_text)
            
            # Add as bullet point
            self.document.add_paragraph(clean_text, style='List Bullet')
            
        except Exception as e:
            print(f"Error adding list item: {e}")
            # Fallback to regular paragraph
            self.document.add_paragraph(text)
    
    def _add_paragraph(self, text):
        """Add regular paragraph"""
        try:
            # Split long text into sentences for better formatting
            sentences = re.split(r'(?<=[.!?])\s+', text)
            
            paragraph = self.document.add_paragraph()
            
            for sentence in sentences:
                sentence = sentence.strip()
                if sentence:
                    # Add some basic formatting detection
                    if sentence.isupper():
                        run = paragraph.add_run(sentence + ' ')
                        run.bold = True
                    else:
                        paragraph.add_run(sentence + ' ')
            
            # Set paragraph formatting
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        except Exception as e:
            print(f"Error adding paragraph: {e}")
            # Simple fallback
            self.document.add_paragraph(text)
    
    def add_metadata(self, title=None, author=None, subject=None):
        """Add metadata to the document"""
        try:
            if title:
                self.document.core_properties.title = title
            if author:
                self.document.core_properties.author = author
            if subject:
                self.document.core_properties.subject = subject
                
        except Exception as e:
            print(f"Error adding metadata: {e}")
    
    def set_page_margins(self, top=1.0, bottom=1.0, left=1.0, right=1.0):
        """Set page margins in inches"""
        try:
            sections = self.document.sections
            for section in sections:
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
                
        except Exception as e:
            print(f"Error setting margins: {e}")
